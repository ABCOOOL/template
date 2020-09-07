from docx import Document

# 创建新docx则不需要添加文件路径，读写存在的docx要加入文件路径
doc = Document()

# 写入内容
doc.add_paragraph("add_content")

# 遍历所有段落
for paragraph in doc.paragraphs:
    print(paragraph.text)

# 保存，新旧文件都需要路径
doc.save("file_path")